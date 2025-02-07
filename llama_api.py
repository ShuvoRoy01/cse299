import pytesseract as tess
from PIL import Image
from docx import Document
import os
import re
import pdfplumber
from PyPDF2 import PdfReader

# Define Tesseract-OCR path
tess.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def extract_text_from_image(image_path):
    # Check if file exists
    if not os.path.exists(image_path):
        print(f"Error: File not found at {image_path}")
        return ""
    
    try:
        # Open the image using PIL
        img = Image.open(image_path)
        
        # Convert image to RGB mode to handle all formats
        img = img.convert("RGB")
        
        # Use pytesseract to extract text with custom config
        custom_config = "--psm 4 -c preserve_interword_spaces=1"
        extracted_text = tess.image_to_string(img, config=custom_config)
        
        return extracted_text.strip()  # Keep text structure, no unnecessary cleaning
    except Exception as e:
        print(f"Error opening image: {e}")
        return ""

def extract_text_from_pdf(pdf_path):
    if not os.path.exists(pdf_path):
        print(f"Error: File not found at {pdf_path}")
        return ""
    
    try:
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(docx_path):
    if not os.path.exists(docx_path):
        print(f"Error: File not found at {docx_path}")
        return ""
    
    try:
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

def save_text_to_docx(text, output_path):
    # Create a new Word document
    doc = Document()
    
    # Add extracted text to the document
    doc.add_paragraph(text)
    
    # Save the document
    doc.save(output_path)

def process_files(folder_path, output_path):
    supported_formats = (".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tiff", ".webp", ".pdf", ".docx")
    doc = Document()
    
    # Check if folder exists
    if not os.path.exists(folder_path):
        print(f"Error: Folder not found at {folder_path}")
        return
    
    # Loop through files in the directory
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        file_extension = os.path.splitext(filename)[1].lower()
        
        if file_extension in supported_formats:
            print(f"Processing {filename}...")
            
            if file_extension in (".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tiff", ".webp"):
                text = extract_text_from_image(file_path)
            elif file_extension == ".pdf":
                text = extract_text_from_pdf(file_path)
            elif file_extension == ".docx":
                text = extract_text_from_docx(file_path)
            else:
                text = ""
            
            if text:
                doc.add_paragraph(f"Extracted text from {filename}:")
                doc.add_paragraph(text)
                doc.add_paragraph("\n--------------------------------\n")
    
    # Save the document if it contains text
    if doc.paragraphs:
        doc.save(output_path)
        print(f"All extracted text saved to {output_path}")
    else:
        print("No text extracted from files.")

def main():
    folder_path = r"E:\\pic to text\\pic"  # Change this to your folder path
    output_path = "output_text.docx"  # Output file path
    
    process_files(folder_path, output_path)

if __name__ == "__main__":
    main()
